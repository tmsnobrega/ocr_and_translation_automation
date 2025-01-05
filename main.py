import os
import re
import shutil
import io
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.auth.transport.requests import Request
from google.oauth2.service_account import Credentials
from PIL import Image
from google.cloud import vision
from google.cloud import translate_v2 as translate
from docx import Document
from docx.shared import RGBColor

# Load environment variables
import creds
FOLDER_ID = creds.FOLDER_ID
GOOGLE_SERVICE_ACCOUNT = creds.GOOGLE_SERVICE_ACCOUNT
LOCAL_IMAGE_FOLDER = creds.LOCAL_IMAGE_FOLDER
LOCAL_WORD_FOLDER = creds.LOCAL_WORD_FOLDER
SCOPES = creds.SCOPES

def authenticate_with_service_account(json_key_path=None, scopes=None):
  """Authenticate using a service account JSON key file."""
  json_key_path = creds.GOOGLE_SERVICE_ACCOUNT

  scopes = creds.SCOPES

  if not os.path.exists(json_key_path):
    raise FileNotFoundError(f"Service account JSON key file not found: {json_key_path}")

  try:
    # Create credentials with the specified scopes
    credentials = Credentials.from_service_account_file(json_key_path, scopes=scopes)
    return credentials
  except Exception as e:
    raise Exception(f"Failed to authenticate with service account: {e}")
    
credentials = authenticate_with_service_account(GOOGLE_SERVICE_ACCOUNT, SCOPES)
service = build('drive', 'v3', credentials=credentials)

# def download_files_from_google_drive():
#   """Download all photo files from Google Drive folder."""
#   print("Downloading files...")
#   query = f"'{FOLDER_ID}' in parents and mimeType contains 'image/'"
#   results = service.files().list(q=query, fields="files(id, name)").execute()
#   items = results.get('files', [])

#   print(f"\tFound {len(items)} files in the folder.")

#   if not os.path.exists(LOCAL_IMAGE_FOLDER):
#     os.makedirs(LOCAL_IMAGE_FOLDER)

#   for item in items:
#     file_path = f"{LOCAL_IMAGE_FOLDER}/{item['name']}"
#     request = service.files().get_media(fileId=item['id'])
#     with open(file_path, 'wb') as f:
#       downloader = MediaIoBaseDownload(f, request)
#       done = False
#       while not done:
#         status, done = downloader.next_chunk()
#         print(f"\tDownloading {item['name']} - {int(status.progress() * 100)}% complete.")

#   print("Download process successfully completed!")

def download_files_from_google_drive():
  """Download all photo files from Google Drive folder."""
  print("Downloading files...")

  # Query to fetch image files in the folder
  query = f"'{FOLDER_ID}' in parents and mimeType contains 'image/'"
  items = []
  page_token = None

  # Fetch all files using pagination
  while True:
    results = service.files().list(
      q=query,
      fields="nextPageToken, files(id, name)",
      pageToken=page_token
    ).execute()
    items.extend(results.get('files', []))
    page_token = results.get('nextPageToken')
    if not page_token:
      break

  # Sort files by name
  items = sorted(items, key=lambda x: x['name'])

  print(f"\tFound {len(items)} files in the folder.")
  if items:
    print(f"\t\tFirst file: {items[0]['name']}")
    print(f"\t\tLast file: {items[-1]['name']}")

  # Ensure local folder exists
  if not os.path.exists(LOCAL_IMAGE_FOLDER):
    os.makedirs(LOCAL_IMAGE_FOLDER)

  # Download each file
  for item in items:
    file_path = f"{LOCAL_IMAGE_FOLDER}/{item['name']}"
    request = service.files().get_media(fileId=item['id'])
    with open(file_path, 'wb') as f:
      downloader = MediaIoBaseDownload(f, request)
      done = False
      while not done:
        status, done = downloader.next_chunk()
        print(f"\tDownloading {item['name']} - {int(status.progress() * 100)}% complete.")

  print("Download process successfully completed!")

def change_photo_names():
  """Rename photo files with the specified naming convention."""
  print("Renaming files...")

  # List files and sort them by name
  files = sorted(os.listdir(LOCAL_IMAGE_FOLDER))

  for idx, file in enumerate(files):
    # Calculate day and photo number based on the index
    day = (idx // 2) + 1
    photo_number = (idx % 2) + 1
    new_name = f"d{day}_p{photo_number}.jpg"
    
    # Perform renaming
    os.rename(f"{LOCAL_IMAGE_FOLDER}/{file}", f"{LOCAL_IMAGE_FOLDER}/{new_name}")
    print(f"\tRenamed {file} to {new_name}")

  print("All files successfully renamed!")

def perform_ocr_with_google_vision():
  """Extract text from photos using OCR via Google Vision API."""
  print("Performing OCR on images...")
  credentials = Credentials.from_service_account_file(GOOGLE_SERVICE_ACCOUNT)
  client = vision.ImageAnnotatorClient(credentials=credentials)

  files = sorted([f for f in os.listdir(LOCAL_IMAGE_FOLDER) if f.endswith('.jpg')])
  ocr_results = {}

  for file in files:
    file_path = f"{LOCAL_IMAGE_FOLDER}/{file}"
    with open(file_path, 'rb') as image_file:
      content = image_file.read()
    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    text = response.full_text_annotation.text if response.full_text_annotation else ""
    ocr_results[file] = text
    print(f"\tOCR completed for {file}")

  print("OCR process successfully completed!")
  return ocr_results
  
def translate_text_to_portuguese(ocr_results):
  """Translate text from English to Portuguese."""
  print("Translating text to Portuguese...")
  credentials = Credentials.from_service_account_file(GOOGLE_SERVICE_ACCOUNT)
  translator = translate.Client(credentials=credentials)
  translated_results = {}

  for file, text in ocr_results.items():
      translated_text = translator.translate(text, source_language='en', target_language='pt')['translatedText']
      translated_results[file] = translated_text
      print(f"\tTranslation completed for {file}")

  print("Translation process successfully completed!")
  return translated_results

def custom_formatting(doc, text, idx):
  """Apply specific formatting rules for the project."""
  text = text.replace("&quot;", "\"")

  # Split the text into sections
  sections = re.split(r"(Reflita sobre a Palavra|Leve para a oração|Mergulhe Mais Fundo)", text)

  # Add the first section as Page header and Header2
  if sections:
    page_header = sections.pop(0).strip()
    if page_header:
      header = doc.add_heading(f"Página {idx}", level=1)
      header.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Change header color to black
      doc.add_heading(page_header, level=2)

  # Process remaining sections
  for section in sections:
    section = section.strip()
    if section in ["Reflita sobre a Palavra", "Leve para a oração", "Mergulhe Mais Fundo"]:
      doc.add_heading(section, level=2)
    elif section:
      # Add line breaks for bullet points
      lines = re.split(r"\n|•", section)
      for line in lines:
        if line.strip():
          doc.add_paragraph(line.strip())

def group_images_by_day(translated_results):
  """Group images by day and save to Word documents."""
  print("Grouping images by day and saving to Word files...")
  grouped_results = {}

  for file, text in translated_results.items():
    day_match = re.match(r'd(\d+)_p\d+.jpg', file)
    if day_match:
      day = f"d{day_match.group(1)}"
      if day not in grouped_results:
        grouped_results[day] = []
      grouped_results[day].append(text)

  if not os.path.exists(LOCAL_WORD_FOLDER):
    os.makedirs(LOCAL_WORD_FOLDER)

  for day, texts in grouped_results.items():
    doc = Document()

    for idx, text in enumerate(texts, 1):
      custom_formatting(doc, text, idx)

    # Save the document
    doc.save(f"{LOCAL_WORD_FOLDER}/{day}.docx")
    print(f"\tSaved {day}.docx")

  print("All Word documents saved successfully!")

def upload_files_to_google_drive():
  """Upload Word documents to Google Drive."""
  print("Uploading Word documents to Google Drive...")
  docs_files = [f for f in os.listdir(LOCAL_WORD_FOLDER) if f.endswith('.docx')]

  for file in docs_files:
    file_metadata = {
      'name': file,
      'parents': [FOLDER_ID]
    }
    media = MediaFileUpload(f"{LOCAL_WORD_FOLDER}/{file}", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print(f"\tUploaded {file}")

  print("All files successfully uploaded to Google Drive!")

def cleanup_temp_files():
  """Delete the temp_files directory and all its contents."""
  print("Cleaning up temporary files...")
  if os.path.exists("./temp_files"):
    shutil.rmtree("./temp_files")
    print("Temporary files deleted successfully!")
  else:
    print("No temporary files to clean up.")

def main():
  download_files_from_google_drive()
  change_photo_names()
  ocr_results = perform_ocr_with_google_vision()
  translated_results = translate_text_to_portuguese(ocr_results)
  group_images_by_day(translated_results)
  upload_files_to_google_drive()
  cleanup_temp_files()
  print("\nAll processes successfully completed!\n")

if __name__ == "__main__":
  main()
